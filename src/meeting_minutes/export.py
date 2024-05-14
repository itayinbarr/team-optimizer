import os
from docx import Document

def save_as_docx(content, filename, directory):
    if not os.path.exists(directory):
        os.makedirs(directory)
    
    file_path = os.path.join(directory, filename)
    doc = Document()
    doc.add_heading("Detailed Tasks for Next Sprint", level=1)
    doc.add_paragraph(content)
    
    doc.save(file_path)

def save_all_tasks_as_docx(tasks, summary, sprint_directory):
    filename = summary.lower().replace(" ", "_") + ".docx"
    save_as_docx(tasks, filename, sprint_directory)

def save_simple_tasks_as_docx(tasks, summary, task_directory):
    simple_tasks = filter_simple_tasks(tasks)
    filename = summary.lower().replace(" ", "_") + "_simple.docx"
    save_as_docx(simple_tasks, filename, task_directory)

def filter_simple_tasks(tasks):
    simple_tasks = ""
    for task in tasks.split('\n\n'):
        if "Complexity level: " in task:
            complexity = int(task.split("Complexity level: ")[1])
            if complexity <= 3:  # Define what "simple enough" means
                simple_tasks += task + "\n\n"
    return simple_tasks
